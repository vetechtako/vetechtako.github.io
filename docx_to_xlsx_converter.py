"""
模擬考題 DOCX -> XLSX 批次轉檔 & 合併工具
適用於 Windows 11 / Python 3.8+
需安裝: pip install python-docx openpyxl
"""

import os
import re
import sys
import traceback
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from pathlib import Path

# ── 檢查套件 ──
_missing = []
try:
    from docx import Document
except ImportError:
    _missing.append("python-docx")

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
except ImportError:
    _missing.append("openpyxl")


# ────────────────────────────────────────
#  DOCX 解析
# ────────────────────────────────────────

def parse_question_cell(text):
    """從題目欄位中解析出題目本文與四個選項。"""
    text = text.strip()
    pattern = r'\(([A-D])\)\s*'
    matches = list(re.finditer(pattern, text))

    if len(matches) < 2:
        pattern2 = r'(?:^|\n)\s*([A-D])\.\s*'
        matches2 = list(re.finditer(pattern2, text))
        if len(matches2) >= 2:
            matches = matches2
            question_text = text[:matches[0].start()].strip()
            options = {}
            for i, m in enumerate(matches):
                label = m.group(1)
                start = m.end()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
                options[label] = text[start:end].strip()
            return question_text, options.get('A', ''), options.get('B', ''), options.get('C', ''), options.get('D', '')

    if len(matches) < 2:
        return text, '', '', '', ''

    question_text = text[:matches[0].start()].strip()
    options = {}
    for i, m in enumerate(matches):
        label = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        options[label] = text[start:end].strip().rstrip('\n').strip()

    return (
        question_text,
        options.get('A', ''),
        options.get('B', ''),
        options.get('C', ''),
        options.get('D', '')
    )


def extract_cell_text(cell):
    """從 docx 表格儲存格中提取完整文字，保留換行。"""
    return '\n'.join(p.text for p in cell.paragraphs)


def convert_docx_to_data(docx_path):
    """讀取 DOCX 檔案中的考題表格，回傳結構化資料。"""
    doc = Document(docx_path)
    all_rows = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header_cells = [extract_cell_text(c).strip() for c in table.rows[0].cells]
        num_cols = len(header_cells)

        col_map = {}
        for i, h in enumerate(header_cells):
            if any(k in h for k in ('題序', '題號', '編號')) or h in ('No', 'no', '#'):
                col_map['num'] = i
            elif any(k in h for k in ('正解', '答案', '解答')):
                col_map['answer'] = i
            elif any(k in h for k in ('題目', '題幹')):
                col_map['question'] = i
            elif '選項A' in h or '選項a' in h:
                col_map['optA'] = i
            elif '選項B' in h or '選項b' in h:
                col_map['optB'] = i
            elif '選項C' in h or '選項c' in h:
                col_map['optC'] = i
            elif '選項D' in h or '選項d' in h:
                col_map['optD'] = i

        if 'question' in col_map and 'answer' in col_map:
            has_opts = all(k in col_map for k in ('optA', 'optB', 'optC', 'optD'))
            for row_idx in range(1, len(table.rows)):
                cells = table.rows[row_idx].cells
                if len(cells) < num_cols:
                    continue
                num_val = extract_cell_text(cells[col_map['num']]).strip() if 'num' in col_map else ''
                answer_val = extract_cell_text(cells[col_map['answer']]).strip()
                if has_opts:
                    q_text = extract_cell_text(cells[col_map['question']]).strip()
                    opt_a = extract_cell_text(cells[col_map['optA']]).strip()
                    opt_b = extract_cell_text(cells[col_map['optB']]).strip()
                    opt_c = extract_cell_text(cells[col_map['optC']]).strip()
                    opt_d = extract_cell_text(cells[col_map['optD']]).strip()
                else:
                    raw_q = extract_cell_text(cells[col_map['question']])
                    q_text, opt_a, opt_b, opt_c, opt_d = parse_question_cell(raw_q)
                if not q_text and not answer_val:
                    continue
                all_rows.append({
                    'num': num_val, 'answer': answer_val, 'question': q_text,
                    'optA': opt_a, 'optB': opt_b, 'optC': opt_c, 'optD': opt_d,
                })
        elif num_cols >= 3:
            for row_idx in range(1, len(table.rows)):
                cells = table.rows[row_idx].cells
                num_val = extract_cell_text(cells[0]).strip()
                answer_val = extract_cell_text(cells[-1]).strip()
                raw_q = extract_cell_text(cells[1])
                q_text, opt_a, opt_b, opt_c, opt_d = parse_question_cell(raw_q)
                if not q_text and not answer_val:
                    continue
                all_rows.append({
                    'num': num_val, 'answer': answer_val, 'question': q_text,
                    'optA': opt_a, 'optB': opt_b, 'optC': opt_c, 'optD': opt_d,
                })

    return all_rows


# ────────────────────────────────────────
#  XLSX 輸出
# ────────────────────────────────────────

HEADERS = ['題號', '答案', '題目', '選項A', '選項B', '選項C', '選項D']
COL_WIDTHS = {'A': 8, 'B': 8, 'C': 45, 'D': 25, 'E': 25, 'F': 25, 'G': 25}


def _apply_header_style(ws, num_data_rows):
    hfont = Font(bold=True, size=11, color='FFFFFF')
    hfill = PatternFill('solid', fgColor='4472C4')
    halign = Alignment(horizontal='center', vertical='center', wrap_text=True)
    border = Border(left=Side(style='thin'), right=Side(style='thin'),
                    top=Side(style='thin'), bottom=Side(style='thin'))
    for col_idx, h in enumerate(HEADERS, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = hfont
        cell.fill = hfill
        cell.alignment = halign
        cell.border = border
    for col_letter, w in COL_WIDTHS.items():
        ws.column_dimensions[col_letter].width = w
    ws.auto_filter.ref = f"A1:G{num_data_rows + 1}"


def save_to_xlsx(data, output_path, sheet_name='Sheet1'):
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    border = Border(left=Side(style='thin'), right=Side(style='thin'),
                    top=Side(style='thin'), bottom=Side(style='thin'))
    wrap = Alignment(vertical='center', wrap_text=True)
    center = Alignment(horizontal='center', vertical='center')

    for row_idx, item in enumerate(data, 2):
        values = [item['num'], item['answer'], item['question'],
                  item['optA'], item['optB'], item['optC'], item['optD']]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = border
            cell.alignment = center if col_idx <= 2 else wrap

    _apply_header_style(ws, len(data))
    wb.save(output_path)
    return len(data)


def merge_xlsx_files(xlsx_paths, output_path, renumber=True, dedup=False):
    wb = Workbook()
    ws = wb.active
    ws.title = '合併結果'

    border = Border(left=Side(style='thin'), right=Side(style='thin'),
                    top=Side(style='thin'), bottom=Side(style='thin'))
    wrap = Alignment(vertical='center', wrap_text=True)
    center = Alignment(horizontal='center', vertical='center')

    current_row = 2
    total_questions = 0
    dup_count = 0
    seen = set()

    for xlsx_path in xlsx_paths:
        src_wb = load_workbook(xlsx_path, data_only=True)
        src_ws = src_wb.active
        for row in src_ws.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
            values = list(row[:7])
            while len(values) < 7:
                values.append('')

            if dedup:
                fingerprint = tuple(str(v).strip() if v is not None else '' for v in values[1:])
                if fingerprint in seen:
                    dup_count += 1
                    continue
                seen.add(fingerprint)

            if renumber:
                values[0] = total_questions + 1

            for col_idx, val in enumerate(values, 1):
                cell = ws.cell(row=current_row, column=col_idx, value=val)
                cell.border = border
                cell.alignment = center if col_idx <= 2 else wrap

            current_row += 1
            total_questions += 1
        src_wb.close()

    _apply_header_style(ws, total_questions)
    wb.save(output_path)
    return total_questions, dup_count


# ────────────────────────────────────────
#  GUI
# ────────────────────────────────────────

def _get_ui_font():
    """取得可用的中文字型。"""
    import tkinter.font as tkfont
    tmp = tk.Tk()
    tmp.withdraw()
    available = set(tkfont.families(tmp))
    tmp.destroy()
    for f in ['Microsoft JhengHei UI', 'Microsoft JhengHei',
              'Microsoft YaHei UI', 'Microsoft YaHei',
              'PMingLiU', 'MingLiU', 'SimSun', 'Arial']:
        if f in available:
            return f
    return 'TkDefaultFont'


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("模擬考題 DOCX -> XLSX 轉檔工具")
        self.geometry("820x680")
        self.minsize(700, 550)

        self._fn = _get_ui_font()

        style = ttk.Style()
        for theme in ('vista', 'winnative', 'default'):
            if theme in style.theme_names():
                style.theme_use(theme)
                break
        style.configure('Title.TLabel', font=(self._fn, 14, 'bold'))
        style.configure('TButton', font=(self._fn, 10), padding=6)
        style.configure('Accent.TButton', font=(self._fn, 10, 'bold'))
        style.configure('TLabelframe.Label', font=(self._fn, 10, 'bold'))
        style.configure('TCheckbutton', font=(self._fn, 10))

        self.docx_files = []
        self.xlsx_files = []
        self.output_dir = None
        self._build_ui()

    def _build_ui(self):
        main = ttk.Frame(self, padding=15)
        main.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main, text="模擬考題 DOCX -> XLSX  批次轉檔 & 合併工具",
                  style='Title.TLabel').pack(pady=(0, 10))

        # ─ 功能一 ─
        f1 = ttk.LabelFrame(main, text="功能一：DOCX -> XLSX 批次轉檔", padding=10)
        f1.pack(fill=tk.X, pady=(0, 10))

        r1 = ttk.Frame(f1)
        r1.pack(fill=tk.X, pady=2)
        ttk.Button(r1, text="選擇 DOCX 檔案 (可多選)", command=self._select_docx).pack(side=tk.LEFT)
        self.docx_count_label = ttk.Label(r1, text="尚未選擇檔案")
        self.docx_count_label.pack(side=tk.LEFT, padx=10)

        self.docx_listbox = tk.Listbox(f1, height=4, font=('Consolas', 9))
        self.docx_listbox.pack(fill=tk.X, pady=5)

        r1b = ttk.Frame(f1)
        r1b.pack(fill=tk.X, pady=2)
        ttk.Button(r1b, text="選擇輸出資料夾", command=self._select_output_dir).pack(side=tk.LEFT)
        self.output_dir_label = ttk.Label(r1b, text="預設：與來源檔案同資料夾")
        self.output_dir_label.pack(side=tk.LEFT, padx=10)

        ttk.Button(f1, text=">>> 開始轉檔 <<<", command=self._run_convert,
                   style='Accent.TButton').pack(pady=5)

        # ─ 功能二 ─
        f2 = ttk.LabelFrame(main, text="功能二：合併多個 XLSX 檔案", padding=10)
        f2.pack(fill=tk.X, pady=(0, 10))

        r2 = ttk.Frame(f2)
        r2.pack(fill=tk.X, pady=2)
        ttk.Button(r2, text="選擇 XLSX 檔案 (可多選)", command=self._select_xlsx).pack(side=tk.LEFT)
        self.xlsx_count_label = ttk.Label(r2, text="尚未選擇檔案")
        self.xlsx_count_label.pack(side=tk.LEFT, padx=10)

        self.xlsx_listbox = tk.Listbox(f2, height=4, font=('Consolas', 9))
        self.xlsx_listbox.pack(fill=tk.X, pady=5)

        opt = ttk.Frame(f2)
        opt.pack(fill=tk.X, pady=2)
        self.renumber_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(opt, text="合併後重新編號題號", variable=self.renumber_var).pack(side=tk.LEFT)
        self.dedup_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(opt, text="移除重複題目", variable=self.dedup_var).pack(side=tk.LEFT, padx=(15, 0))

        ttk.Button(f2, text=">>> 開始合併 <<<", command=self._run_merge,
                   style='Accent.TButton').pack(pady=5)

        # ─ 日誌 ─
        lf = ttk.LabelFrame(main, text="執行日誌", padding=5)
        lf.pack(fill=tk.BOTH, expand=True)
        self.log_text = scrolledtext.ScrolledText(lf, height=8, font=('Consolas', 9),
                                                   state=tk.DISABLED, wrap=tk.WORD)
        self.log_text.pack(fill=tk.BOTH, expand=True)

        self.status_var = tk.StringVar(value="就緒")
        ttk.Label(main, textvariable=self.status_var,
                  font=(self._fn, 9), foreground='gray').pack(anchor=tk.W, pady=(5, 0))

    def _log(self, msg):
        self.log_text.configure(state=tk.NORMAL)
        self.log_text.insert(tk.END, msg + '\n')
        self.log_text.see(tk.END)
        self.log_text.configure(state=tk.DISABLED)

    def _select_docx(self):
        files = filedialog.askopenfilenames(
            title="選擇 DOCX 檔案",
            filetypes=[("Word 文件", "*.docx"), ("所有檔案", "*.*")])
        if files:
            self.docx_files = list(files)
            self.docx_listbox.delete(0, tk.END)
            for f in self.docx_files:
                self.docx_listbox.insert(tk.END, os.path.basename(f))
            self.docx_count_label.config(text="已選擇 {} 個檔案".format(len(self.docx_files)))

    def _select_output_dir(self):
        d = filedialog.askdirectory(title="選擇輸出資料夾")
        if d:
            self.output_dir = d
            self.output_dir_label.config(text=d)

    def _select_xlsx(self):
        files = filedialog.askopenfilenames(
            title="選擇 XLSX 檔案",
            filetypes=[("Excel 檔案", "*.xlsx"), ("所有檔案", "*.*")])
        if files:
            self.xlsx_files = list(files)
            self.xlsx_listbox.delete(0, tk.END)
            for f in self.xlsx_files:
                self.xlsx_listbox.insert(tk.END, os.path.basename(f))
            self.xlsx_count_label.config(text="已選擇 {} 個檔案".format(len(self.xlsx_files)))

    def _run_convert(self):
        if not self.docx_files:
            messagebox.showwarning("提示", "請先選擇 DOCX 檔案！")
            return
        self.status_var.set("轉檔中，請稍候...")
        threading.Thread(target=self._convert_worker, daemon=True).start()

    def _convert_worker(self):
        success, fail = 0, 0
        output_files = []
        for docx_path in self.docx_files:
            basename = Path(docx_path).stem
            out_dir = self.output_dir or os.path.dirname(docx_path)
            out_path = os.path.join(out_dir, "{}.xlsx".format(basename))
            try:
                self.after(0, self._log, "[轉檔] {}".format(os.path.basename(docx_path)))
                data = convert_docx_to_data(docx_path)
                if not data:
                    self.after(0, self._log, "  [警告] 未偵測到考題表格，跳過。")
                    fail += 1
                    continue
                count = save_to_xlsx(data, out_path, sheet_name=basename[:31])
                self.after(0, self._log, "  [OK] 共 {} 題 -> {}".format(count, os.path.basename(out_path)))
                output_files.append(out_path)
                success += 1
            except Exception as e:
                self.after(0, self._log, "  [失敗] {}".format(e))
                fail += 1

        summary = "轉檔完成：成功 {} 個，失敗 {} 個".format(success, fail)
        self.after(0, self._log, "\n{}\n{}\n".format('=' * 50, summary))
        self.after(0, self.status_var.set, summary)
        if output_files:
            self.after(0, self._log, "[提示] 轉出的 XLSX 已自動加入合併檔案清單中。")
            self.after(0, self._auto_add_xlsx, output_files)

    def _auto_add_xlsx(self, files):
        self.xlsx_files = files
        self.xlsx_listbox.delete(0, tk.END)
        for f in files:
            self.xlsx_listbox.insert(tk.END, os.path.basename(f))
        self.xlsx_count_label.config(text="已選擇 {} 個檔案 (自動)".format(len(files)))

    def _run_merge(self):
        if not self.xlsx_files:
            messagebox.showwarning("提示", "請先選擇 XLSX 檔案！")
            return
        if len(self.xlsx_files) < 2:
            messagebox.showwarning("提示", "至少需要 2 個 XLSX 檔案才能合併！")
            return
        out_path = filedialog.asksaveasfilename(
            title="儲存合併結果", defaultextension=".xlsx",
            filetypes=[("Excel 檔案", "*.xlsx")],
            initialfile="合併考題.xlsx")
        if not out_path:
            return
        self.status_var.set("合併中，請稍候...")
        threading.Thread(target=self._merge_worker, args=(out_path,), daemon=True).start()

    def _merge_worker(self, out_path):
        try:
            self.after(0, self._log, "\n[合併] 開始合併 {} 個檔案...".format(len(self.xlsx_files)))
            for f in self.xlsx_files:
                self.after(0, self._log, "  - {}".format(os.path.basename(f)))
            total, dup_count = merge_xlsx_files(
                self.xlsx_files, out_path,
                renumber=self.renumber_var.get(),
                dedup=self.dedup_var.get())
            dup_msg = "（移除 {} 題重複）".format(dup_count) if dup_count > 0 else ""
            msg = "[OK] 合併完成！共 {} 題{} -> {}".format(total, dup_msg, os.path.basename(out_path))
            self.after(0, self._log, msg)
            self.after(0, self.status_var.set, msg)
            self.after(0, lambda: messagebox.showinfo(
                "完成", "合併完成！\n共 {} 題{}\n儲存至：{}".format(total, dup_msg, out_path)))
        except Exception as e:
            self.after(0, self._log, "[失敗] 合併失敗：{}".format(e))
            self.after(0, self.status_var.set, "合併失敗")


# ────────────────────────────────────────
#  啟動入口
# ────────────────────────────────────────

def main():
    if _missing:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror(
            "缺少套件",
            "請先安裝以下 Python 套件再執行：\n\n"
            "  pip install {}\n\n"
            "缺少：{}".format(' '.join(_missing), ', '.join(_missing)))
        root.destroy()
        sys.exit(1)

    try:
        app = App()
        app.mainloop()
    except Exception:
        error_msg = traceback.format_exc()
        try:
            root = tk.Tk()
            root.withdraw()
            messagebox.showerror("啟動錯誤", "程式啟動失敗：\n\n{}".format(error_msg))
            root.destroy()
        except Exception:
            print("程式啟動失敗：\n{}".format(error_msg), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
